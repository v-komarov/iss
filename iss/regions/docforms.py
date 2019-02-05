#coding:utf-8

import StringIO
import datetime

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

"""
 Сборка функций формирования печатных форм 


"""




def monthname(n):
    """Преобразование номера месяца в название"""

    mon = {
        1: u"Января",
        2: u"Февраля",
        3: u"Марта",
        4: u"Апреля",
        5: u"Мая",
        6: u"Июня",
        7: u"Июля",
        8: u"Августа",
        9: u"Сентября",
        10: u"Октября",
        11: u"Ноября",
        12: u"Декабря",
    }

    return mon[n]





# Печатная форма для акта АВР
def doc1(avr_obj):


    f = StringIO.StringIO()

    document = Document()

    ### УТВЕРЖДАЮ
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Pt(250)
    p.add_run(u'УТВЕРЖДАЮ:\nТехнический директор\nАО «СибТрансТелеКом»\nА.П. Ляднов\n\n\n"____" _____________ 20 _____ г.\n').bold = True

    ### Заголовок
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'Акт АВР № %s\nорганизации и выполнения АВР и расследования причин аварии корпоративного заказа' % avr_obj.id).bold = True

    ### АО и дата
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'АО «СибТрансТелеКом»                                    «{}» {} {} г.\n'.format(avr_obj.datetime_avr.day, monthname(avr_obj.datetime_avr.month), avr_obj.datetime_avr.year)).bold = True


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Место проведения работ: {} {}".format(avr_obj.city.name, avr_obj.address, avr_obj.objnet)).font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Участок: {}".format(avr_obj.area)).font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Пусковой комплекс: {}".format(avr_obj.complex)).font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Номер ЛР\ЛРП и наименование (суть) работ: «{}» {} {} г.\n".format(avr_obj.datetime_avr.day, monthname(avr_obj.datetime_avr.month), avr_obj.datetime_avr.year)).font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Мы, нижеподписавшиеся:").font.size = Pt(9)

    ### Состав комиссии , должности
    if avr_obj.commission:
        pos = avr_obj.commission.position.split(";")
        for po in pos:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(u"{}".format(po)).font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"составили настоящий акт в том, что по результатам анализа причин и характера повреждения, произошедшего «{}» {} {} г.\nна участке {}\nвыявлено следующее:".format(avr_obj.datetime_avr.day, monthname(avr_obj.datetime_avr.month), avr_obj.datetime_avr.year,avr_obj.area)).font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1. Хронология аварийно-восстановительных работ:\n1.1. Регистрация аварии:\n").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"предварительная диагностика типа аварийной ситуации и локализация повреждения:\n").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.2. Регистрация отказа: в ___ час ___ мин").font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.3. Оповещение  ______: в ___ час ___ мин    Сбор в  ___ час ___ мин").font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.4. Определено место повреждения\nна расстоянии __________ метров от оптического кросса _____________").font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ## дата выезда
    date_out = u"" if avr_obj.datetime_work == None else u"«{}» {} {}".format(avr_obj.datetime_work.day,monthname(avr_obj.datetime_work.month),avr_obj.datetime_work.year)
    p.add_run(u"1.5. Выезд на место повреждения {} в ____ час ____ мин  Прибытие ____ час _____ мин".format(date_out)).font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.6. На месте повреждения\nОбнаружено:\nпо причине:\n").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.7. Восстановительные работы по временной схеме не проводились\nначаты в _____ час _____ мин ___________ 2018г\nзаконченыв _____ час _____ мин ___________ 2018г").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Потери на восстановление связи по временной схеме составили ________ час _______ мин.\n c ______ по _______").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Способ временного восстановления связи:\nПричины задержек в восстановлении связи:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.8. Восстановительные работы по постоянной схеме составили: ____ часов ____ мин\nc ____ час _____ мин      _____________ 2018г.\nпо ____ час _____ мин      _____________ 2018г.").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"2. Перечень материалов, израсходованных при проведении АВР:").font.size = Pt(9)



    nn = avr_obj.store_out_set.count() ### Количество строк с материалом


    table = document.add_table(1+nn,6, style="TableGrid")
    table.columns[0].width = Pt(40)
    table.columns[1].width = Pt(200)
    table.columns[2].width = Pt(30)
    table.columns[3].width = Pt(40)
    table.columns[4].width = Pt(100)
    table.columns[5].width = Pt(60)
    table.autofit = False
    heading_cells = table.rows[0].cells
    heading_cells[0].text = u"№пп"
    heading_cells[1].text = u"Наименование материала"
    heading_cells[2].text = u"Ед.\nиз."
    heading_cells[3].text = u"Расход"
    heading_cells[4].text = u"Название склада"
    heading_cells[5].text = u"Из состава\nоборудо-\nвания"
    heading_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[1].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[2].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[3].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[4].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[5].paragraphs[0].runs[0].font.size = Pt(7)


    n = 1
    for row in avr_obj.store_out_set.all():
        cells = table.rows[n].cells
        cells[0].text = "%s" % n
        cells[1].text = row.store_rest.name
        cells[2].text = row.store_rest.dimension
        cells[3].text = "%.2f" % row.q
        cells[4].text = row.store_rest.store.name
        cells[5].text = ""
        cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        cells[1].paragraphs[0].runs[0].font.size = Pt(7)
        cells[2].paragraphs[0].runs[0].font.size = Pt(7)
        cells[3].paragraphs[0].runs[0].font.size = Pt(7)
        cells[4].paragraphs[0].runs[0].font.size = Pt(7)
        cells[5].paragraphs[0].runs[0].font.size = Pt(7)

        n += 1

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"\n3. Выводы комиссии:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"3.1. Причины повреждения:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"3.2. Ответственные за повреждения, принятые меры:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"3.3. Оценка организации аварийно-восстановительных работ (указать, при наличии таковых, причины задержки восстановления действия связи):").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"3.4. Принятые меры по оптимизации аварийно-восстановительных работ и профилактике повреждений:").font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"\n4. Подписи членов комиссии:").font.size = Pt(9)

    ### Подписи комиссии
    if avr_obj.commission:
        sign = avr_obj.commission.sign.split(";")
        for s in sign:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(u"__________________________________________ {}".format(s)).font.size = Pt(9)


    document.save(f)

    data = f.getvalue()

    return data





# Печатная форма для акта регламентных работ
def doc2(avr_obj):


    f = StringIO.StringIO()

    document = Document()

    ### УТВЕРЖДАЮ
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Pt(250)
    p.add_run(u'УТВЕРЖДАЮ:\nТехнический директор\nАО «СибТрансТелеКом»\nА.П. Ляднов\n\n\n"____" _____________ 20 _____ г.\n').bold = True

    ### Заголовок
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'Акт АВР № %s\nорганизации и выполнения регламентных работ' % avr_obj.id).bold = True

    ### АО и дата
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(u'АО «СибТрансТелеКом»                                    «{}» {} {} г.\n'.format(avr_obj.datetime_avr.day, monthname(avr_obj.datetime_avr.month), avr_obj.datetime_avr.year)).bold = True


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Место проведения работ: Адрес: {} {}".format(avr_obj.city.name, avr_obj.address, avr_obj.objnet)).font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Участок: {}".format(avr_obj.area)).font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Пусковой комплекс: {}".format(avr_obj.complex)).font.size = Pt(9)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Номер ЛР\ЛРП и наименование (суть) работ: «{}» {} {} г.\n".format(avr_obj.datetime_avr.day, monthname(avr_obj.datetime_avr.month), avr_obj.datetime_avr.year)).font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Мы, нижеподписавшиеся:").font.size = Pt(9)

    ### Состав комиссии , должности
    if avr_obj.commission:
        pos = avr_obj.commission.position.split(";")
        for po in pos:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(u"{}".format(po)).font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"составили настоящий акт в том, что по результатам анализа причин и характера работ, произошедших «{}» {} {} г.\nвыявлено следующее: ______________________________________________________________".format(avr_obj.datetime_avr.day, monthname(avr_obj.datetime_avr.month), avr_obj.datetime_avr.year)).font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1. Хронология аварийно-восстановительных работ").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"1.2 Регламентные работы по постоянной схеме начаты:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(u"в ___ час ___ мин  «___» _________ 201 __ г.").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"Окончены:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(u"в ___ час ___ мин  «___» _________ 201 __ г.").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"2. Перечень материалов, израсходованных при проведении АВР:").font.size = Pt(9)



    nn = avr_obj.store_out_set.count() ### Количество строк с материалом


    table = document.add_table(1+nn,6, style="TableGrid")
    table.columns[0].width = Pt(40)
    table.columns[1].width = Pt(200)
    table.columns[2].width = Pt(30)
    table.columns[3].width = Pt(40)
    table.columns[4].width = Pt(100)
    table.columns[5].width = Pt(60)
    table.autofit = False
    heading_cells = table.rows[0].cells
    heading_cells[0].text = u"№пп"
    heading_cells[1].text = u"Наименование материала"
    heading_cells[2].text = u"Ед.\nиз."
    heading_cells[3].text = u"Расход"
    heading_cells[4].text = u"Название склада"
    heading_cells[5].text = u"Из состава\nоборудо-\nвания"
    heading_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[1].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[2].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[3].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[4].paragraphs[0].runs[0].font.size = Pt(7)
    heading_cells[5].paragraphs[0].runs[0].font.size = Pt(7)


    n = 1
    for row in avr_obj.store_out_set.all():
        cells = table.rows[n].cells
        cells[0].text = "%s" % n
        cells[1].text = row.store_rest.name
        cells[2].text = row.store_rest.dimension
        cells[3].text = "%.2f" % row.q
        cells[4].text = row.store_rest.store.name
        cells[5].text = ""
        cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        cells[1].paragraphs[0].runs[0].font.size = Pt(7)
        cells[2].paragraphs[0].runs[0].font.size = Pt(7)
        cells[3].paragraphs[0].runs[0].font.size = Pt(7)
        cells[4].paragraphs[0].runs[0].font.size = Pt(7)
        cells[5].paragraphs[0].runs[0].font.size = Pt(7)

        n += 1

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"\n3. Выполнены работы:").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"3.1. Причины работ: ___________________________________________").font.size = Pt(9)

    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"3.2. Выводы комиссии: ______________________________________________").font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"\n4. Подписи членов комиссии:").font.size = Pt(9)

    ### Подписи комиссии
    if avr_obj.commission:
        sign = avr_obj.commission.sign.split(";")
        for s in sign:
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(u"__________________________________________ {}".format(s)).font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"\n4. Материальные ценности отнесены на соответствующую статью учета").font.size = Pt(9)


    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run(u"\n5. Бухгалтер материалист _______________________ «___» _________ 201 __ г.").font.size = Pt(9)



    document.save(f)

    data = f.getvalue()

    return data